from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT_FILE = ROOT / "College_Management_System_End_User_Manual.docx"
HERO_IMAGE = ROOT / "frontend" / "src" / "assets" / "hero.png"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_borders(section):
    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is None:
        pg_borders = OxmlElement("w:pgBorders")
        sect_pr.append(pg_borders)

    for edge in ("top", "left", "bottom", "right"):
        element = pg_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            pg_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "12")
        element.set(qn("w:color"), "D6E6F2")


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    set_page_borders(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for style_name in ("Title", "Subtitle"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(15, 76, 129)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(29, 53, 87)


def add_paragraph(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F9FF")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.italic = True
    run.font.color.rgb = RGBColor(65, 90, 119)
    doc.add_paragraph()


def add_numbered_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "1F4E78")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(255, 255, 255)

    for row_index, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            if row_index % 2 == 0:
                set_cell_shading(cell, "F7FAFC")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(value)
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)

    doc.add_paragraph()


def build_doc():
    doc = Document()
    style_document(doc)

    add_paragraph(
        doc,
        "College Management System",
        bold=True,
        size=22,
        color=(23, 58, 94),
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        doc,
        "End User Manual",
        bold=True,
        size=20,
        color=(23, 58, 94),
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        doc,
        "Prepared from the current project structure and role-based pages in the CampusOne / EduManager web application.",
        size=11,
        color=(75, 93, 114),
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        doc,
        f"Document date: {date.today().strftime('%d %B %Y')}",
        size=11,
        color=(75, 93, 114),
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    if HERO_IMAGE.exists():
        doc.add_picture(str(HERO_IMAGE), width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_callout(
        doc,
        "This end user manual explains how a user can open the website, sign in, move through the menu, and complete the main tasks available for public visitors, students, faculty, HOD, and admin users.",
    )

    doc.add_heading("1. Website Access", level=1)
    add_numbered_list(
        doc,
        [
            "Open the website in your browser. The public landing page shows the college home screen with the main navigation and action buttons like Apply Now and Login.",
            "If you are a new visitor, start from the public pages: Home, About, Contact, and Register.",
            "If you already have an account, open the Login page.",
            "After login, the system automatically redirects you to the correct dashboard based on your role.",
        ],
    )

    doc.add_heading("2. Demo Login Credentials", level=1)
    add_paragraph(
        doc,
        "Use the following seeded demo accounts when the sample database has been loaded:",
    )
    add_table(
        doc,
        ["Role", "Email", "Password", "Result after login"],
        [
            ["Admin", "admin@college.edu", "admin123", "Opens /admin/dashboard"],
            ["HOD", "hod@college.edu", "faculty123", "Opens /hod/dashboard"],
            ["Faculty", "faculty1@college.edu", "faculty123", "Opens /faculty/dashboard"],
            ["Student", "student1@college.edu", "student123", "Opens /student/dashboard"],
        ],
    )
    add_callout(
        doc,
        "Important note: On the login screen, the user first selects a role. Faculty users must also choose either HOD Login or Faculty Login before submitting email and password.",
    )

    doc.add_heading("3. Role Overview", level=1)
    add_paragraph(
        doc,
        "The table below summarizes what each role is expected to do inside the website.",
    )
    add_table(
        doc,
        ["Role", "Main purpose in website", "What this role can do"],
        [
            [
                "Public Visitor",
                "Access public information and create an account",
                "View home, about, and contact pages; open register page; create a new account; open login page",
            ],
            [
                "Student",
                "Use self-service academic and campus facilities",
                "View dashboard, attendance, timetable, results, fees, notifications, placements, library, certificates, hostel, transport, materials, complaints, service requests, request status, and documents",
            ],
            [
                "Faculty",
                "Manage teaching and classroom academic work",
                "View dashboard, profile, timetable, notifications; mark attendance; use face attendance; enter marks; upload results; manage assignments; upload study materials; review students; request leave; view payslips",
            ],
            [
                "HOD",
                "Manage department-level academic supervision",
                "Open HOD console; review department faculty and department overview; manage courses and timetable; approve leave; review department notifications",
            ],
            [
                "Admin",
                "Control institution-wide academic and administrative modules",
                "Manage students, faculty, departments, courses, admissions, attendance, exams, results, fees, timetable, notifications, library, placements, hostel, transport, payroll, scholarships, events, certificates, reports, and approvals",
            ],
        ],
    )
    add_callout(
        doc,
        "The website is role-based, so each user will only see the menu items and workflows relevant to that role after login.",
    )

    doc.add_heading("4. Public Visitor Flow", level=1)
    doc.add_heading("A. Home Page", level=2)
    add_numbered_list(
        doc,
        [
            "Open the landing page to view the college introduction, feature cards, statistics, testimonials, and footer contact information.",
            "Use Apply Now to move to the registration page.",
            "Use Login if you already have a student, faculty, HOD, or admin account.",
        ],
    )
    doc.add_heading("B. Register Page", level=2)
    add_numbered_list(
        doc,
        [
            "Open Register from the navbar or the home page.",
            "Enter Full name, Email, and Password.",
            "Choose the Role from Student, Faculty, or Admin.",
            "If you choose Student, enter the Roll number field.",
            "If you choose Faculty, enter the Employee ID field.",
            "Click Register to create the account.",
            "After successful registration, the website redirects the user directly to the role dashboard.",
        ],
    )

    doc.add_heading("5. Login Flow", level=1)
    add_numbered_list(
        doc,
        [
            "Open the Login page.",
            "Select one of the role cards: Admin, Faculty, or Student.",
            "If the selected role is Faculty, choose the subtype: HOD Login or Faculty Login.",
            "Confirm that the email and password fields are filled correctly.",
            "Click Sign In.",
            "The system redirects the user to the correct portal: student dashboard, faculty dashboard, HOD dashboard, or admin dashboard.",
        ],
    )

    doc.add_heading("6. Student User Guide", level=1)
    add_paragraph(
        doc,
        "Student users mainly work from Academic Path, Learning Lab, Campus Life, and Support Desk sections.",
    )
    doc.add_heading("A. What a Student Can Do", level=2)
    add_bullet_list(
        doc,
        [
            "See academic performance from the dashboard.",
            "Track attendance and view timetable details.",
            "Check results and fee status.",
            "Read notifications and upcoming exam information.",
            "Access placements, certificates, library, hostel, and transport modules when available.",
            "Upload documents, register face ID, submit complaints, and create service requests.",
        ],
    )
    doc.add_heading("B. Student Dashboard", level=2)
    add_numbered_list(
        doc,
        [
            "Review the dashboard cards for Attendance, Current CGPA, Pending Fees, and Active Courses.",
            "Use quick actions to open View Timetable, Check Results, Pay Fees, or Apply Placement.",
            "Read recent notifications and upcoming exams.",
            "Check the attendance overview graph to monitor subject-wise attendance percentage.",
        ],
    )
    doc.add_heading("C. Student Day-to-Day Tasks", level=2)
    add_bullet_list(
        doc,
        [
            "Profile: View or update personal information.",
            "Attendance: Check attendance details and status by subject.",
            "Timetable: View the class schedule.",
            "Results: Open published examination results.",
            "Fees: Check dues and complete fee payment.",
            "Notifications: Read college notices and updates.",
            "Placements: Review drives and apply for eligible opportunities.",
            "Library: View issued books or library records.",
            "Certificates: View generated certificates and download them when available.",
            "Hostel and Transport: Access service-related details if assigned.",
        ],
    )
    doc.add_heading("D. Student Support Desk Flow", level=2)
    add_numbered_list(
        doc,
        [
            "Open Face ID Register to register facial identity if that feature is enabled.",
            "Open Complaints to submit a grievance.",
            "Open Service Request to submit a request for campus support.",
            "Use Request Status to track submitted requests.",
            "Use Documents to upload or manage student documents.",
        ],
    )

    doc.add_heading("7. Faculty User Guide", level=1)
    add_paragraph(
        doc,
        "Faculty users mainly operate from Operations, Academic Flow, and Governance sections in the sidebar.",
    )
    doc.add_heading("A. What a Faculty User Can Do", level=2)
    add_bullet_list(
        doc,
        [
            "Monitor teaching-related information from the faculty dashboard.",
            "Mark attendance manually or with face attendance tools.",
            "Enter marks, upload results, and manage assignments.",
            "Upload study materials for students.",
            "View student lists, notifications, leave requests, and payslips.",
        ],
    )
    doc.add_heading("B. Faculty Dashboard", level=2)
    add_numbered_list(
        doc,
        [
            "Sign in using a faculty account.",
            "Open the dashboard to review teaching-related data and shortcuts.",
            "Use the sidebar to switch quickly between attendance, marks, result upload, materials, and notifications.",
        ],
    )
    doc.add_heading("C. Mark Attendance", level=2)
    add_numbered_list(
        doc,
        [
            "Open Mark Attendance from the sidebar.",
            "Choose the assigned Course from the dropdown.",
            "Select the Date.",
            "Load the student list automatically for the selected course.",
            "Use All present or All absent for quick entry if needed.",
            "For each student, mark one status: Present, Absent, or Late.",
            "Click Submit attendance to save.",
            "Review the Attendance report section below the form to verify totals, sessions, and percentages.",
        ],
    )
    doc.add_heading("D. Other Faculty Tasks", level=2)
    add_bullet_list(
        doc,
        [
            "Face Attendance: Capture or manage face-based attendance workflows.",
            "Enter Marks: Enter marks for exams or internal assessments.",
            "Assignments: Create and manage assignment records.",
            "Result Upload: Upload result data for publication workflows.",
            "Study Materials: Upload or manage LMS or study content.",
            "My Students: View assigned student data.",
            "Notifications: Review institutional messages.",
            "Leave Request: Submit leave applications.",
            "My Payslips: Review payroll and salary records.",
        ],
    )

    doc.add_heading("8. HOD User Guide", level=1)
    add_paragraph(
        doc,
        "HOD users sign in through the faculty role but land in a department management portal.",
    )
    doc.add_heading("A. What an HOD Can Do", level=2)
    add_bullet_list(
        doc,
        [
            "Monitor department-level operations from the HOD console.",
            "Review department faculty and department overview pages.",
            "Manage courses and department timetable information.",
            "Approve or reject leave requests.",
            "Review department notifications and academic coordination tasks.",
        ],
    )
    doc.add_heading("B. HOD Workflow", level=2)
    add_numbered_list(
        doc,
        [
            "On the login page, choose Faculty and then HOD Login.",
            "After sign-in, open the HOD Console dashboard.",
            "Use Dept Faculty to manage or review faculty members in the department.",
            "Use Dept Overview to check the department view of students and operations.",
            "Open Courses to manage course-related work.",
            "Open Timetable to manage department schedule details.",
            "Open Leave Approvals to approve or reject faculty leave requests.",
            "Use Notifications to review department communication.",
        ],
    )

    doc.add_heading("9. Admin User Guide", level=1)
    add_paragraph(
        doc,
        "Admin users have the broadest menu. They control institution-wide academic and administrative modules.",
    )
    doc.add_heading("A. What an Admin Can Do", level=2)
    add_bullet_list(
        doc,
        [
            "View institution-wide dashboard metrics and analytics.",
            "Create, update, filter, export, and remove master records such as students, faculty, departments, and courses.",
            "Control admissions, attendance, exams, results, fees, timetable, and notifications.",
            "Manage library, placements, hostel, transport, payroll, scholarship, event, certificate, and report modules.",
            "Perform approval and governance tasks across the system.",
        ],
    )
    doc.add_heading("B. Admin Dashboard", level=2)
    add_numbered_list(
        doc,
        [
            "Sign in as Admin and open the dashboard.",
            "Review summary cards for Total Students, Total Faculty, Departments, Courses, Pending Admissions, and Fee Collection.",
            "Read the charts for admission trends, fee collection overview, department-wise students, and attendance by department.",
            "Use the quick actions panel to jump to Add student, Add faculty, Create notice, or View reports.",
        ],
    )
    doc.add_heading("C. Student Management Workflow", level=2)
    add_numbered_list(
        doc,
        [
            "Open Students from the sidebar.",
            "Use the search field to find a student by roll number, name, or email.",
            "Use the filters for Department, Semester, and Batch.",
            "Click Add student to open the form.",
            "Fill in the required data such as roll number, full name, email, password, department, semester, and status.",
            "Optionally add personal information, address details, and guardian details.",
            "Click Create student to save the new record.",
            "To edit a student, open the row and save after changes.",
            "To remove a student, use the delete action and confirm the dialog.",
            "Use Excel or CSV export to download the filtered list.",
        ],
    )
    doc.add_heading("D. Admin Core Operations", level=2)
    add_bullet_list(
        doc,
        [
            "Faculty: Add, edit, and manage faculty records.",
            "Departments: Maintain department master data.",
            "Curriculum: Manage courses and academic structures.",
            "Attendance: Review institution attendance records.",
            "Exams and Results: Manage examinations and published outcomes.",
            "Timetable and Result Upload: Control scheduling and uploads.",
            "Hostel, Transport, Library, and Placements: Maintain operational service modules.",
            "Fees and Payroll: Handle finance-related records.",
            "Admissions: Review and process applications.",
            "Notifications: Publish notices to users.",
            "Leave Approvals: Manage approval workflow.",
            "Reports: View analytics and summaries.",
            "Scholarship, Events, and Certificates: Manage specialized modules.",
        ],
    )

    doc.add_heading("10. Recommended End-to-End Usage Sequence", level=1)
    add_numbered_list(
        doc,
        [
            "Start at the public website and review college information.",
            "Create an account from Register if the user does not already exist.",
            "Sign in through the role-based login page.",
            "Use the dashboard first to understand pending tasks and summary status.",
            "Use the sidebar to open the correct module for the next action.",
            "Complete data entry or review work and confirm that the updated tables or messages show the save succeeded.",
            "Return to the dashboard or relevant reports page to verify the effect of the action.",
            "Sign out when finished so the next user starts with a clean session.",
        ],
    )

    doc.add_heading("11. Troubleshooting Tips", level=1)
    add_bullet_list(
        doc,
        [
            "If login fails, check that the role selected on the login page matches the account type.",
            "If HOD login does not open the HOD portal, use the HOD sub-role option instead of regular faculty login.",
            "If data pages show placeholders or empty cards, verify that the backend API and database are running.",
            "If a table appears empty, try clearing filters or refreshing the page.",
            "If uploads do not work, confirm that the selected file matches the expected type.",
            "If pages do not load after login, verify that the frontend and backend are connected and the seeded users exist.",
        ],
    )

    add_paragraph(
        doc,
        "This end user manual was generated from the current website routes, sidebar structure, login flow, and representative page behavior in the project source code.",
        italic=True,
        size=10,
        color=(75, 93, 114),
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.save(str(OUTPUT_FILE))
    print(f"Generated: {OUTPUT_FILE}")


if __name__ == "__main__":
    build_doc()
